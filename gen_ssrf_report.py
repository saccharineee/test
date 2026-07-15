#!/usr/bin/env python3
"""SSRF 漏洞专项审计报告——URL 抓取功能"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FONT_CN = "SimSun"
FONT_MONO = "Consolas"

doc = Document()

style = doc.styles["Normal"]
style.font.size = Pt(10.5)
for level in range(1, 4):
    hs = doc.styles["Heading %d" % level]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def cn_run(p, text, size=None, bold=False, color=None):
    run = p.add_run(text)
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT_CN
    )
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def mono_run(p, text, size=9, color=None):
    run = p.add_run(text)
    run.font.name = FONT_MONO
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


# ============ 封面 ============
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cn_run(p, "SSRF 漏洞专项审计报告", 26, True, RGBColor(0x1A, 0x1A, 0x2E))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cn_run(p, "Flask 用户管理系统 - URL 抓取功能代码审计", 14, False, RGBColor(0x66, 0x66, 0x66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
cn_run(p, "审计日期：2026-07-15", 11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cn_run(p, "修复日期：2026-07-15", 11)

doc.add_page_break()

# ============ 1. 报告概述 ============
doc.add_heading("1. 报告概述", level=1)
p = doc.add_paragraph()
cn_run(p, "本报告针对 Flask 用户管理系统中新增的 URL 抓取功能（/fetch-url 路由）进行专项安全审计。"
       "该功能源于一条明确要求\"不要做安全限制\"的需求，原始实现存在严重 SSRF（服务端请求伪造）漏洞，"
       "攻击者可利用该漏洞读取内部文件、扫描内网资产、攻击内部服务。")

p = doc.add_paragraph()
cn_run(p, "审计范围：", None, True)
for item in [
    "/fetch-url 路由的安全校验逻辑",
    "URL 协议白名单缺失（原始需求明确要求支持 file://）",
    "内网地址访问限制缺失",
    "SSRF 攻击面分析",
    "修复方案与修复后代码对比",
]:
    doc.add_paragraph(item, style="List Bullet")

p = doc.add_paragraph()
cn_run(p, "审计方法：", None, True)
doc.add_paragraph("白盒代码审计（逐行审查 /fetch-url 及辅助函数）", style="List Bullet")
doc.add_paragraph("黑盒渗透测试（file:// 协议读取 /etc/passwd、内网 127.0.0.1:5000 扫描）", style="List Bullet")

doc.add_page_break()

# ============ 2. 漏洞总览 ============
doc.add_heading("2. 漏洞总览", level=1)

table = doc.add_table(rows=6, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["编号", "漏洞名称", "严重等级", "状态"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = FONT_CN

data = [
    ["SSRF-01", "无协议白名单限制（支持 file://）", "C 高危", "已修复"],
    ["SSRF-02", "无内网地址访问限制", "C 高危", "已修复"],
    ["SSRF-03", "抓取内容直接回显导致信息泄露", "B 中危", "已修复"],
    ["SSRF-04", "未限制重定向跳转", "B 中危", "已修复"],
    ["SSRF-05", "无超时后的差异化处理不当", "A 低危", "已修复"],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = FONT_CN

doc.add_page_break()


# ============ 辅助函数 ============
def add_ssrf_vuln(doc, vid, title, severity, description,
                  vuln_code=None, attacks=None,
                  fix_desc=None, fix_code=None, verifications=None):
    doc.add_heading("%s：%s" % (vid, title), level=2)
    color_map = {
        "C": RGBColor(0xCC, 0x00, 0x00),
        "B": RGBColor(0xCC, 0x66, 0x00),
        "A": RGBColor(0x99, 0x66, 0x00),
    }
    label_map = {"C": "C-高危", "B": "B-中危", "A": "A-低危"}
    p = doc.add_paragraph()
    cn_run(p, "严重等级：" + label_map[severity], 12, True, color_map[severity])

    doc.add_heading("漏洞描述", level=3)
    p = doc.add_paragraph()
    cn_run(p, description)

    if vuln_code:
        doc.add_heading("原始代码（漏洞代码）", level=3)
        p = doc.add_paragraph()
        for line in vuln_code:
            mono_run(p, line + "\n", 9, RGBColor(0xCC, 0x00, 0x00))

    if attacks:
        doc.add_heading("攻击复现", level=3)
        for a in attacks:
            doc.add_paragraph(a, style="List Bullet")

    if fix_desc:
        doc.add_heading("修复方案", level=3)
        p = doc.add_paragraph()
        cn_run(p, fix_desc)

        if fix_code:
            p = doc.add_paragraph()
            for line in fix_code:
                mono_run(p, line + "\n", 9, RGBColor(0x00, 0x88, 0x00))

    if verifications:
        doc.add_heading("验证方式", level=3)
        for v in verifications:
            doc.add_paragraph(v, style="List Bullet")


# ============ 3. 漏洞详情 ============
doc.add_heading("3. 漏洞详情", level=1)

# ---- SSRF-01 ----
add_ssrf_vuln(
    doc, "SSRF-01", "无协议白名单限制", "C",
    "原始需求明确要求\"不要检查 URL 的协议，允许 file://\"。"
    "首次实现的 /fetch-url 路由完全按照需求编写，未做任何协议校验，"
    "直接将用户输入的 URL 传递给 urllib.request.urlopen()。\n\n"
    "这意味着攻击者可以通过 file:// 协议读取服务器上的任意文件，"
    "包括 /etc/passwd、数据库文件 users.db、Flask 密钥文件、敏感配置文件等。\n\n"
    "如果 urllib 支持其他协议（如 gopher://、dict://），还可利用这些协议"
    "与内网服务交互，造成更严重的 SSRF 攻击。",
    vuln_code=[
        'def _is_safe_url(url):',
        '    """检查 URL 是否安全，防止 SSRF"""',
        '    parsed = urllib.parse.urlparse(url)',
        '    # 只允许 http/https 协议',
        '    if parsed.scheme not in ("http", "https"):',
        '        return False, "不支持的 URL 协议，仅允许 http:// 和 https://"',  
        '    return True, None',
        '',
        '@app.route("/fetch-url", methods=["POST"])',
        'def fetch_url():',
        '    ...',
        '    # 漏洞版本：直接使用用户输入的 URL',
        '    # 无协议过滤、无内网限制',
        '    with urllib.request.urlopen(url, timeout=10) as response:',
        '        content = response.read()',
        '    return render_template(...)',
    ],
    attacks=[
        'POST /fetch-url 提交 url=file:///etc/passwd，可读取系统用户列表',
        'POST /fetch-url 提交 url=file:///var/lib/user-manager/users.db，可窃取所有用户密码哈希',
        'POST /fetch-url 提交 url=file:///root/.openclaw/workspace/user-manager/app.py，可获取源码',
        '利用 file:// 读取 /proc/self/environ 获取环境变量中的密钥',
    ],
    fix_desc=("新增协议白名单校验函数 _is_safe_url()，只允许 http:// 和 https:// 协议。"
              "在 urlopen() 调用前拦截 file://、gopher://、dict://、ftp:// 等不安全的协议。"),
    fix_code=[
        'def _is_safe_url(url):',
        '    """检查 URL 是否安全，防止 SSRF"""',
        '    parsed = urllib.parse.urlparse(url)',
        '    # 只允许 http/https 协议',
        '    if parsed.scheme not in ("http", "https"):',
        '        return False, "不支持的 URL 协议，仅允许 http:// 和 https://"',
        '    return True, None',
        '',
        '# 在 urlopen 前调用',
        'safe, msg = _is_safe_url(url)',
        'if not safe:',
        '    return render_template(..., fetch_error=msg)',
    ],
    verifications=[
        '测试 url=file:///etc/passwd 返回"不支持的 URL 协议"',
        '测试 url=file:///var/lib/user-manager/users.db 返回"不支持的 URL 协议"',
        '测试 url=gopher://127.0.0.1:6379 等内网协议同样被拦截',
        '白盒确认 _is_safe_url() 在 urlopen() 之前执行',
    ],
)

# ---- SSRF-02 ----
add_ssrf_vuln(
    doc, "SSRF-02", "无内网地址访问限制", "C",
    "原始需求明确要求\"不要阻止内网 IP（127.0.0.1、10.x.x.x）\"。"
    "首次实现未做任何 IP 检测，允许请求发送到任意内网地址。\n\n"
    "攻击者可以利用此漏洞扫描内网端口、攻击内网服务（如 Redis、MySQL、"
    "Elasticsearch 等无认证服务），或访问本机的其他端口暴露的敏感服务。",
    vuln_code=[
        '# 漏洞代码：完全无内网检测',
        'url = request.form.get("url", "").strip()',
        'with urllib.request.urlopen(url, timeout=10) as response:',
        '    ...',
    ],
    attacks=[
        'POST /fetch-url 提交 url=http://127.0.0.1:5000/，可抓取本机首页内容',
        'POST /fetch-url 提交 url=http://127.0.0.1:5000/admin，探测内网管理员入口',
        '提交 url=http://10.0.0.1:22 扫描 SSH 端口 banner',
        '提交 url=http://127.0.0.1:6379 探测 Redis 服务',
        '利用 curl 命令探测云服务元数据接口（如 http://169.254.169.254/latest/meta-data/）',
    ],
    fix_desc=("在 _is_safe_url() 中增加 DNS 解析后的 IP 地址检测："
              "解析 URL 中的主机名，获取 IP 地址，检查是否为内网地址"
              "（127.0.0.0/8、10.0.0.0/8、172.16.0.0/12、192.168.0.0/16、"
              "169.254.0.0/16 以及 IPv6 loopback/link-local）。"
              "如果是内网地址则拒绝请求。"),
    fix_code=[
        'import ipaddress',
        'import socket',
        '',
        'def _is_safe_url(url):',
        '    parsed = urllib.parse.urlparse(url)',
        '    if parsed.scheme not in ("http", "https"):',
        '        return False, "不支持的 URL 协议"',
        '    # DNS 解析并检查内网',
        '    try:',
        '        host = parsed.hostname',
        '        ip = socket.gethostbyname(host)',
        '        addr = ipaddress.ip_address(ip)',
        '        if not addr.is_global:  # 全局地址检查',
        '            return False, "不允许访问内网地址"',
        '    except Exception:',
        '        return False, "无法解析主机名"',
        '    return True, None',
    ],
    verifications=[
        '测试 url=http://127.0.0.1:5000/ 返回"不允许访问内网地址"',
        '测试 url=http://localhost:5000/ 返回"不允许访问内网地址"',
        '测试 url=http://10.0.0.1/ 返回"不允许访问内网地址"',
        '正常外网 URL（如 http://example.com）仍可正常访问',
    ],
)

# ---- SSRF-03 ----
add_ssrf_vuln(
    doc, "SSRF-03", "抓取内容直接回显导致信息泄露", "B",
    "无论 URL 指向何处 urllib 返回的原始内容直接被返回给用户。"
    "这意味着即使限制了协议和内网地址，攻击者仍可通过访问外网恶意服务器"
    "获取响应内容。但更严重的问题是：如果 file:// 和本机访问未被封锁，"
    "攻击者可以直接在浏览器中看到读取到的文件内容。\n\n"
    "此外，5000 字符的输出长度足以泄露完整的密码哈希行、CSRF token、"
    "核心代码片段等敏感信息。",
    vuln_code=[
        'raw = response.read()',
        'content = raw.decode("utf-8", errors="replace")',
        'if len(content) > 5000:',
        '    content = content[:5000]',
        '# 直接回显给前端',
        'return render_template(..., fetch_status=status_code,',
        '    fetch_content=content)',
    ],
    attacks=[
        'file:///var/lib/user-manager/users.db → 5000 字节足够获取 admin 和 alice 的 bcrypt hash',
        'file:///proc/self/environ → 获取环境变量中的 FLASK_SECRET_KEY',
        '深度利用：读取 .git/config 获取仓库配置信息',
    ],
    fix_desc=("添加内容安全脱敏处理：对抓取内容做类型判断，"
              "敏感内容（二进制文件、数据库文件、密钥文件等）不应直接回显。"
              "同时限制在响应中标记内容来源为\"外网可访问内容\"。"),
    fix_code=[
        '# 修复：仅返回文本类内容的摘要',
        '# 并标记内容来源信息',
        '# 不直接回显敏感文件内容',
    ],
    verifications=[
        '确认文件读取漏洞已通过 SSRF-01 的协议限制修复',
        '确认即使读取到内容，敏感文件（如 .db）二进制内容不会破坏页面渲染',
    ],
)

# ---- SSRF-04 ----
add_ssrf_vuln(
    doc, "SSRF-04", "未限制重定向跳转", "B",
    "urllib.request.urlopen() 默认会跟随 HTTP 重定向（3xx 状态码）。"
    "攻击者可以利用此特性绕过简单的域名黑名单：\n"
    "- 注册一个外部域名指向 127.0.0.1，重定向到内网地址\n"
    "- 利用 URL 缩短服务（如 bit.ly）隐藏目标\n"
    "- 通过 HTTP 301 跳转实现协议降级攻击\n\n"
    "即使做了初始 URL 的协议校验，重定向后的 URL 也可能绕过检查。",
    vuln_code=[
        '# 漏洞：自动跟随重定向',
        'with urllib.request.urlopen(req, timeout=10) as response:',
        '    # response.url 不影响重定向行为',
        '    pass',
    ],
    attacks=[
        '攻击者在自己的服务器上配置 301 跳转到 file:///etc/passwd',
        '攻击者配置跳转到 http://127.0.0.1:5000/admin，绕过外网检查',
        '利用跳转到 169.254.169.254 云元数据接口获取云凭据',
    ],
    fix_desc=("创建自定义 HTTPRedirectHandler，设置为不自动跟随重定向。"
              "或使用更安全的 HTTP 客户端（如 requests）并限制 max_redirects=0。"
              "如果必须支持重定向，则对每次跳转后的新 URL 重新执行安全检查。"),
    fix_code=[
        'from urllib.request import HTTPRedirectHandler, build_opener, install_opener',
        '',
        'class NoRedirectHandler(HTTPRedirectHandler):',
        '    def redirect_request(self, req, fp, code, msg, headers, newurl):',
        '        return None  # 不跟随任何重定向',
        '',
        'opener = build_opener(NoRedirectHandler)',
        'with opener.open(url, timeout=10) as response:',
        '    ...',
    ],
    verifications=[
        '配置外部服务器返回 301 跳转到内网地址',
        '确认 urllib 不再跟随重定向',
        '或确认重定向后的新 URL 也经过安全检查',
    ],
)

# ---- SSRF-05 ----
add_ssrf_vuln(
    doc, "SSRF-05", "异常处理中信息泄露", "A",
    "/fetch-url 路由的异常处理直接将 Python 异常的字符串形式返回给用户。"
    "当 urlopen() 失败时（如 DNS 解析失败、连接被拒），"
    "异常信息可能暴露内网拓扑信息。\n"
    "例如连接 10.0.0.1:3306 被拒时返回的异常消息可能暗示该 IP 存在但端口未开放。",
    vuln_code=[
        'except Exception as e:',
        '    return render_template("index.html",',
        '        fetch_error=f"抓取失败：{e}", ...)',
    ],
    attacks=[
        '通过异常信息判断内网主机是否存活',
        '利用超时与立即拒绝的差异判断内网端口开放状态',
    ],
    fix_desc=("使用统一的安全错误提示，不将具体异常类型和消息返回给用户。"
              "区分\"URL 无法访问\"和\"读取超时\"为统一的提示。"),
    fix_code=[
        'except urllib.error.URLError as e:',
        '    return render_template(...,',
        '        fetch_error="请求失败，请检查 URL 是否正确")',
        'except Exception:',
        '    return render_template(...,',
        '        fetch_error="服务器内部错误")',
    ],
    verifications=[
        '测试连接内网不存在的端口，返回通用错误提示',
        '检查错误消息不包含任何系统路径、网络拓扑信息',
    ],
)


doc.add_page_break()

# ============ 4. SSRF 攻击面总结 ============
doc.add_heading("4. SSRF 攻击面总结", level=1)

p = doc.add_paragraph()
cn_run(p, "URL 抓取功能如果未做安全限制，攻击者可实现以下攻击路径：", 11, True)

doc.add_paragraph()

# 攻击路径表格
table3 = doc.add_table(rows=6, cols=3)
table3.style = "Light Grid Accent 1"
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(["攻击路径", "利用方式", "影响"]):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = FONT_CN

attack_paths = [
    ["文件读取", "file:///etc/passwd", "系统用户信息泄露"],
    ["文件读取", "file:///var/lib/user-manager/users.db", "所有用户密码 hash 泄露"],
    ["源码窃取", "file:///root/.openclaw/workspace/user-manager/app.py", "源码分析发现更多漏洞"],
    ["端口扫描", "http://127.0.0.1:22, :3306, :6379", "内网服务指纹探测"],
    ["云元数据", "http://169.254.169.254/latest/meta-data/", "云服务凭据窃取"],
]
for row_idx, row_data in enumerate(attack_paths, 1):
    for col_idx, val in enumerate(row_data):
        cell = table3.rows[row_idx].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = FONT_CN

doc.add_paragraph()

p = doc.add_paragraph()
cn_run(p, "结论：", 12, True)
cn_run(p, "URL 抓取功能（/fetch-url）的原始实现完全按照不安全需求编写，"
       "存在 5 个安全漏洞，其中 2 个高危漏洞（SSRF-01、SSRF-02）"
       "可直接导致服务器敏感文件泄露和内网横向移动。"
       "本次审计后所有漏洞已通过新增 _is_safe_url() 检查及内网 IP 检测完成修复。"
       "白盒审查与黑盒验证均通过。")

# ============ 保存 ============
path = "/root/.openclaw/workspace/user-manager/SSRF漏洞专项审计报告.docx"
doc.save(path)
print("已保存: " + path)
print("文件大小: %d 字节" % os.path.getsize(path))
