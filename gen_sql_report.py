#!/usr/bin/env python3
"""生成 SQL 专项漏洞审计报告，覆盖原有 docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ============ 样式 ============
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_severity(doc, text, level):
    colors = {
        '🔴': RGBColor(0xCC, 0x00, 0x00),
        '🟠': RGBColor(0xCC, 0x66, 0x00),
        '🟡': RGBColor(0x99, 0x66, 0x00),
    }
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = colors.get(level, RGBColor(0,0,0))

# ============ 封面 ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run('SQL 专项漏洞审计报告')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Flask 用户管理系统 — 基于黑盒渗透测试')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
run = p.add_run('审计日期：2026-07-08')
run.font.size = Pt(11)

doc.add_page_break()

# ============ 1. 报告概述 ============
doc.add_heading('1. 报告概述', level=1)
doc.add_paragraph(
    '本报告针对 Flask 用户管理系统中与 SQL 相关的所有安全问题进行专项审计。'
    '涵盖 SQL 注入、数据库架构安全、凭证存储安全、'
    'SQL 层面的暴力破解防护等维度。'
)

doc.add_paragraph('审计范围：')
items = [
    'SQL 注入漏洞（Register、Search、Login 接口）',
    '数据库架构安全性（多数据库、数据分裂）',
    '密码/凭证存储方式',
    'SQL 层面的速率限制与账户锁定',
    '数据库访问控制与权限',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('审计方法：')
run.bold = True
doc.add_paragraph('黑盒渗透测试（模拟攻击者视角，无源代码依赖）', style='List Bullet')
doc.add_paragraph('白盒代码审计（验证修复措施完整性）', style='List Bullet')

doc.add_page_break()

# ============ 2. 漏洞总览 ============
doc.add_heading('2. 漏洞总览', level=1)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['编号', '漏洞名称', '严重等级', '状态']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['SQL-01', 'SQL 注入（Register 接口）', '🔴 高危', '✅ 已修复'],
    ['SQL-02', 'SQL 注入（Search 接口）', '🔴 高危', '✅ 已修复'],
    ['SQL-03', '双数据库并行架构', '🟠 中危', '✅ 已修复'],
    ['SQL-04', '明文密码存储', '🟠 中危', '✅ 已修复'],
    ['SQL-05', 'SQL 层暴力破解防护缺失', '🟡 低危', '✅ 已修复'],
    ['SQL-06', 'Search 接口 SQL 信息泄露', '🟡 低危', '✅ 已修复'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_page_break()

# ============ 3. 各漏洞详情 ============
doc.add_heading('3. 漏洞详情', level=1)

# ---- SQL-01 ----
doc.add_heading('3.1 SQL-01：SQL 注入（Register 接口）', level=2)
add_severity(doc, '严重等级：🔴 高危', '🔴')

doc.add_heading('漏洞描述', level=3)
doc.add_paragraph(
    '用户注册接口使用 Python f-string 直接拼接用户输入构建 SQL INSERT 语句，'
    '攻击者可在用户名、密码等字段中注入恶意 SQL 代码，'
    '导致任意 SQL 语句执行。'
)

doc.add_heading('原始代码（漏洞代码）', level=3)
p = doc.add_paragraph()
run = p.add_run(
    'sql = f"INSERT INTO users (username, password, email, phone) '
    "VALUES ('{username}', '{password}', '{email}', '{phone}')" + '"\n'
    'ndb.execute(sql)'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_heading('攻击复现', level=3)
doc.add_paragraph(
    '输入用户名：admin\' --\n'
    '拼接后 SQL：INSERT INTO users ... VALUES (\'admin\' --\', ...)\n'
    '效果：注释掉后续 SQL，绕过所有约束。', style='List Bullet'
)
doc.add_paragraph(
    '输入用户名：\', \'hacked\', \'x@x.com\', \'123\'); --\n'
    '效果：插入任意数据，篡改数据库。', style='List Bullet'
)

doc.add_heading('修复方案', level=3)
doc.add_paragraph(
    '全部改用参数化查询（prepared statement），使用 ? 占位符传参，'
    '确保用户输入被数据库引擎视为数据而非代码。'
)

doc.add_heading('修复后代码', level=3)
p = doc.add_paragraph()
run = p.add_run(
    'db.execute(\n'
    '    "INSERT INTO users (username, password_hash, email, phone) '
    'VALUES (?, ?, ?, ?)",\n'
    '    (username, password_hash, email, phone),\n'
    ')'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

doc.add_heading('验证方式', level=3)
doc.add_paragraph('黑盒测试：发送注入 payload 到 /register，确认 SQL 注入失败，用户输入被安全存储为字符串。')
doc.add_paragraph('白盒审计：确认所有数据库查询均使用 ? 占位符，无 f-string 拼接。')

# ---- SQL-02 ----
doc.add_heading('3.2 SQL-02：SQL 注入（Search 接口）', level=2)
add_severity(doc, '严重等级：🔴 高危', '🔴')

doc.add_heading('漏洞描述', level=3)
doc.add_paragraph(
    '用户搜索接口使用 f-string 拼接关键字构造 LIKE 查询，'
    '攻击者可通过搜索框注入 SQL 代码，'
    '导出全量用户数据或执行任意 SQL。'
)

doc.add_heading('原始代码（漏洞代码）', level=3)
p = doc.add_paragraph()
run = p.add_run(
    'sql = f"SELECT * FROM users WHERE '
    "username LIKE '%{keyword}%' OR email LIKE '%{keyword}%'" + '"\n'
    'rows = ndb.execute(sql).fetchall()'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_heading('攻击复现', level=3)
doc.add_paragraph(
    '搜索：\' OR 1=1 --\n'
    '拼接后：SELECT * FROM users WHERE '
    "username LIKE '%' OR 1=1 --%'\n"
    '效果：返回所有用户记录。', style='List Bullet'
)
doc.add_paragraph(
    '搜索：\' UNION SELECT id, username, password_hash, email FROM users --\n'
    '效果：通过 UNION 注入泄露其他表字段或敏感列。', style='List Bullet'
)
doc.add_paragraph(
    '搜索：\'; DROP TABLE users; --\n'
    '效果：删除用户表，导致系统不可用。', style='List Bullet'
)

doc.add_heading('修复方案', level=3)
doc.add_paragraph('改用参数化查询，LIKE 模式也作为参数传入：')

p = doc.add_paragraph()
run = p.add_run(
    'like_pattern = f"%{keyword}%"\n'
    'rows = db.execute(\n'
    '    "SELECT * FROM users WHERE username LIKE ? OR email LIKE ?",\n'
    '    (like_pattern, like_pattern),\n'
    ').fetchall()'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

doc.add_heading('验证方式', level=3)
doc.add_paragraph('黑盒测试：发送 \' OR 1=1 -- 等 payload，确认只返回实际匹配结果，无数据泄露。')
doc.add_paragraph('白盒审计：确认 LIKE 子句也使用参数化查询。')

# ---- SQL-03 ----
doc.add_heading('3.3 SQL-03：双数据库并行架构', level=2)
add_severity(doc, '严重等级：🟠 中危', '🟠')

doc.add_heading('漏洞描述', level=3)
doc.add_paragraph(
    '系统维护了两套独立的 SQLite 数据库：'
)
doc.add_paragraph('users.db（主库） — 用于登录验证，bcrypt 哈希密码，参数化查询。', style='List Bullet')
doc.add_paragraph('data/users.db（教学库） — 用于注册和搜索，明文密码，f-string 拼接 SQL。', style='List Bullet')

doc.add_heading('业务影响', level=3)
doc.add_paragraph('注册功能不可用：注册写教学库，登录查主库，注册的用户无法登录。', style='List Bullet')
doc.add_paragraph('数据分裂：两套用户表独立维护，一致性无法保证。', style='List Bullet')
doc.add_paragraph('攻击面翻倍：多一套数据库系统，多一次配置和维护风险。', style='List Bullet')

doc.add_heading('修复方案', level=3)
doc.add_paragraph('移除 data/users.db 教学库，统一使用单一主库 users.db。注册和搜索均走主库且使用参数化查询。')
doc.add_paragraph('原教学库文件已备份为 data/users.db.legacy.bak。')

doc.add_heading('验证方式', level=3)
doc.add_paragraph('确认 data/users.db 已不存在（或已备份）。')
doc.add_paragraph('注册新用户后直接登录，验证功能正常。')

# ---- SQL-04 ----
doc.add_heading('3.4 SQL-04：明文密码存储', level=2)
add_severity(doc, '严重等级：🟠 中危', '🟠')

doc.add_heading('漏洞描述', level=3)
doc.add_paragraph(
    '教学库 data/users.db 以明文存储密码（字段名为 password），'
    '且包含预置的测试用户（admin/admin123、alice/alice2025）。'
    '一旦数据库被注入获取数据，密码直接暴露。'
)

doc.add_heading('修复方案', level=3)
doc.add_paragraph('统一到主库后，全部使用 bcrypt(password + 随机盐) 哈希存储。')
doc.add_paragraph('密码哈希字段名为 password_hash，明确语义。')

doc.add_heading('验证方式', level=3)
doc.add_paragraph('SQLite 直接查询数据库：SELECT password_hash FROM users，确认返回值为 bcrypt 哈希串（以 $2b$ 开头）。')
doc.add_paragraph('确认无任何位置存储明文密码。')

# ---- SQL-05 ----
doc.add_heading('3.5 SQL-05：SQL 层暴力破解防护缺失', level=2)
add_severity(doc, '严重等级：🟡 低危', '🟡')

doc.add_heading('漏洞描述', level=3)
doc.add_paragraph(
    '原始版本无速率限制和账户锁定机制。登录失败次数仅存于进程内存中，'
    '服务重启即重置。攻击者可无限次尝试密码。'
)

doc.add_heading('修复方案', level=3)
doc.add_paragraph('SQL 层面：在 users 表中增加 login_attempts 和 locked_until 字段。', style='List Bullet')
doc.add_paragraph('持久化速率记录：写入 data/rate_state.json，重启不丢失。', style='List Bullet')
doc.add_paragraph('每 IP 15 分钟内最多 20 次登录尝试。', style='List Bullet')
doc.add_paragraph('单账户 5 次失败后锁定 30 分钟。', style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('关键 SQL 语句：')
run.bold = True
p = doc.add_paragraph()
run = p.add_run(
    '-- 尝试次数递增\n'
    'UPDATE users SET login_attempts = login_attempts + 1 WHERE username = ?\n\n'
    '-- 锁定期写入\n'
    'UPDATE users SET locked_until = ? WHERE username = ?\n\n'
    '-- 登录成功后清零\n'
    'UPDATE users SET login_attempts = 0, locked_until = NULL WHERE username = ?'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('验证方式', level=3)
doc.add_paragraph('黑盒测试：连续 20 次 POST /login 后第 21 次被限速。')
doc.add_paragraph('数据库查询：确认 login_attempts 和 locked_until 字段正常更新。')
doc.add_paragraph('重启测试：确认速率计数器重启后仍然有效。')

# ---- SQL-06 ----
doc.add_heading('3.6 SQL-06：Search 接口信息泄露', level=2)
add_severity(doc, '严重等级：🟡 低危', '🟡')

doc.add_heading('漏洞描述', level=3)
doc.add_paragraph(
    'Search 接口未要求登录认证，任意访问者可通过 SQL 搜索枚举所有用户信息'
    '（用户名、邮箱、手机号）。结合 SQL-02 的注入漏洞，攻击者可获取完整用户数据库。'
)

doc.add_heading('修复方案', level=3)
doc.add_paragraph('Search 路由增加登录检测：')

p = doc.add_paragraph()
run = p.add_run(
    '@app.route("/search")\n'
    'def search():\n'
    '    username = session.get("username")\n'
    '    if not username:\n'
    '        return redirect("/login")\n'
    '    ...'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

doc.add_heading('验证方式', level=3)
doc.add_paragraph('不带 Cookie 访问 /search → 302 跳转 /login。')
doc.add_paragraph('带上有效登录 Cookie → 正常返回搜索结果。')

doc.add_page_break()

# ============ 4. 最终验证结果 ============
doc.add_heading('4. 最终黑盒验证结果', level=1)

table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['编号', '验证项', '结果']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

verify_data = [
    ['SQL-01', "Register SQL注入: ' OR 1=1 --", '✅ 被参数化查询拦截'],
    ['SQL-01', 'Register SQL注入: DROP TABLE', '✅ 被参数化查询拦截'],
    ['SQL-02', "Search SQL注入: ' UNION SELECT", '✅ 被参数化查询拦截'],
    ['SQL-03', '双数据库已移除', '✅ 统一为单一主库'],
    ['SQL-04', '密码存储方式', '✅ bcrypt 哈希'],
    ['SQL-05', '速率限制（20次/15min）+ 账户锁定', '✅ 第21次触发限速，5次触发锁定'],
]
for row_idx, row_data in enumerate(verify_data, 1):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph()

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('结论：')
run.bold = True
run.font.size = Pt(12)
p.add_run(' 所有 SQL 相关漏洞已全部修复。黑盒测试未发现可被利用的 SQL 安全漏洞。')

# ============ 保存 ============
path = '/root/.openclaw/workspace/user-manager/完整漏洞审计报告.docx'
doc.save(path)
print(f'✅ 已保存: {path}')
print(f'   文件大小: {os.path.getsize(path)} 字节')
